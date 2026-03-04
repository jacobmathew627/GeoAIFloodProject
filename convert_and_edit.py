import os
import re
from pdf2docx import Converter
from docx import Document

pdf_file = "Phase1Report_Fixed.pdf"
docx_file = "Phase1Report_Fixed.docx"

print("Converting PDF to DOCX...")
cv = Converter(pdf_file)
cv.convert(docx_file, start=0, end=None)
cv.close()

print("Opening DOCX for text replacement...")
doc = Document(docx_file)

def replace_text(text):
    return re.sub(r'(?i)lstm', 'CNN', text)

replacements = 0

# Replace in paragraphs
for p in doc.paragraphs:
    for run in p.runs:
        if re.search(r'(?i)lstm', run.text):
            run.text = replace_text(run.text)
            replacements += 1

# Replace in tables
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    if re.search(r'(?i)lstm', run.text):
                        run.text = replace_text(run.text)
                        replacements += 1

doc.save(docx_file)
print(f"Done! Replaced {replacements} occurrences of 'lstm'. Saved to {docx_file}")

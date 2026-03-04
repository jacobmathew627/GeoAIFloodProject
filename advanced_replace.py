import re
from docx import Document

docx_file = "Phase1Report_Fixed.docx"
output_file = "Phase1Report_CNN.docx"

print("Opening DOCX for advanced text replacement...")
doc = Document(docx_file)

# Ordered dictionary of replacements (longer phrases first to avoid partial matches later)
replacements_map = {
    # Full names and phrases
    r'(?i)Long Short-?Term Memory \(?LSTM\)?': 'Convolutional Neural Network (CNN)',
    r'(?i)Long Short-?Term Memory': 'Convolutional Neural Network',
    
    # Concept descriptions
    r'(?i)LSTM-based rainfall prediction': 'CNN-based flood susceptibility mapping',
    r'(?i)LSTM-based': 'CNN-based',
    r'(?i)predict short-term rainfall trends based on historical rainfall time-series': 'map flood susceptibility based on multi-variate spatial data',
    r'(?i)analyze sequential data like rainfall time-series': 'analyze spatial data like topography and hydrology',
    r'(?i)analyze historical rainfall patterns': 'analyze spatial risk factors',
    r'(?i)historical rainfall data spanning multiple years': 'spatial data representing risk factors',
    r'(?i)historical rainfall time-series': 'spatial risk factors',
    r'(?i)historical rainfall data': 'spatial risk data',
    
    # Specific LSTM architecture terms
    r'(?i)temporal dependencies': 'spatial dependencies',
    r'(?i)time-series and': 'spatial and',
    r'(?i)time-series': 'spatial data',
    r'(?i)Rainfall Risk Score': 'Flood Susceptibility Score',
    r'(?i)memory cells': 'convolutional filters',
    r'(?i)Recurrent Neural Networks \(RNNs\)': 'Convolutional Neural Networks (CNNs)',
    r'(?i)Recurrent Neural Networks': 'Convolutional Neural Networks',
    r'(?i)gated mechanism': 'encoder-decoder structure',
    r'(?i)cumulative rainfall effects': 'spatial feature representations',
    r'(?i)delayed and': 'complex and',
    r'(?i)vanishing gradient problem': 'spatial resolution loss',

    # Model references
    r'(?i)LSTM model': 'U-Net CNN model',
    r'(?i)LSTM network': 'U-Net CNN',
    r'(?i)LSTM layer': 'CNN layer',
    r'(?i)LSTM unit': 'CNN filter',
    r'(?i)LSTM cell': 'CNN building block',

    # Catch-all
    r'(?i)\blstm\b': 'U-Net CNN'
}

stats = {k: 0 for k in replacements_map.keys()}

def replace_text_advanced(text):
    original_text = text
    for pattern, replacement in replacements_map.items():
        new_text, count = re.subn(pattern, replacement, text)
        if count > 0:
            stats[pattern] += count
            text = new_text
    return text

# Replace in paragraphs
for p in doc.paragraphs:
    for run in p.runs:
        if run.text.strip():
            run.text = replace_text_advanced(run.text)

# Replace in tables
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    if run.text.strip():
                        run.text = replace_text_advanced(run.text)

doc.save(output_file)

print(f"Saved optimized document to {output_file}")
print("\nReplacement Statistics:")
for pattern, count in stats.items():
    if count > 0:
        print(f"'{pattern}' -> {count} replacements")
